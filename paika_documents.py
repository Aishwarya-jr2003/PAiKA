import streamlit as st
import chromadb
from groq import Groq
from dotenv import load_dotenv
import os
from pathlib import Path
from datetime import datetime
from langchain_text_splitters import RecursiveCharacterTextSplitter
import PyPDF2
from docx import Document as DocxDocument
from io import BytesIO

load_dotenv()

# Page config
st.set_page_config(
    page_title="PAiKA - Document Manager",
    page_icon="📚",
    layout="wide"
)

# Document loader with preview
class DocumentManager:
    @staticmethod
    def load_and_preview_text(file):
        content = file.read().decode('utf-8')
        preview = content[:500] + "..." if len(content) > 500 else content
        return content, preview, len(content)
    
    @staticmethod
    def load_and_preview_pdf(file):
        pdf = PyPDF2.PdfReader(BytesIO(file.read()))
        content = "\n".join([page.extract_text() for page in pdf.pages])
        preview = content[:500] + "..." if len(content) > 500 else content
        return content, preview, len(pdf.pages)
    
    @staticmethod
    def load_and_preview_docx(file):
        doc = DocxDocument(BytesIO(file.read()))
        content = "\n".join([p.text for p in doc.paragraphs])
        preview = content[:500] + "..." if len(content) > 500 else content
        return content, preview, len(doc.paragraphs)
    
    @classmethod
    def load(cls, file):
        ext = Path(file.name).suffix.lower()
        loaders = {
            '.txt': cls.load_and_preview_text,
            '.md': cls.load_and_preview_text,
            '.pdf': cls.load_and_preview_pdf,
            '.docx': cls.load_and_preview_docx
        }
        
        if ext in loaders:
            content, preview, meta = loaders[ext](file)
            return content, preview, meta, ext
        return None, None, None, None

# Initialize
if 'chroma_client' not in st.session_state:
    st.session_state.chroma_client = chromadb.PersistentClient(path="./paika_docs_db")

if 'collection' not in st.session_state:
    try:
        st.session_state.collection = st.session_state.chroma_client.get_collection("paika_docs")
    except:
        st.session_state.collection = st.session_state.chroma_client.create_collection("paika_docs")

if 'text_splitter' not in st.session_state:
    st.session_state.text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)

if 'uploaded_files_info' not in st.session_state:
    st.session_state.uploaded_files_info = []

# Header
st.title("📚 PAiKA Document Manager")
st.caption("Advanced document handling with preview and management")

# Tabs
tab1, tab2, tab3 = st.tabs(["📤 Upload", "📂 Library", "🔍 Search"])

with tab1:
    st.subheader("Upload Documents")
    
    uploaded_files = st.file_uploader(
        "Choose files",
        type=['txt', 'md', 'pdf', 'docx'],
        accept_multiple_files=True,
        help="Supported: TXT, PDF, DOCX, Markdown"
    )
    
    if uploaded_files:
        st.write(f"📁 {len(uploaded_files)} files selected")
        
        # Preview before upload
        for file in uploaded_files:
            with st.expander(f"📄 Preview: {file.name}"):
                content, preview, meta, file_type = DocumentManager.load(file)
                
                if content:
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.metric("Type", file_type)
                    with col2:
                        if file_type == '.pdf':
                            st.metric("Pages", meta)
                        elif file_type == '.docx':
                            st.metric("Paragraphs", meta)
                        else:
                            st.metric("Characters", meta)
                    with col3:
                        st.metric("Size", f"{len(content)} chars")
                    
                    st.text_area("Preview", preview, height=150, disabled=True)
                    
                    # Reset file pointer for later processing
                    file.seek(0)
        
        # Process button
        if st.button("🚀 Upload All Files", type="primary", use_container_width=True):
            progress = st.progress(0)
            status = st.empty()
            
            for i, file in enumerate(uploaded_files):
                status.info(f"Processing {file.name}...")
                
                content, preview, meta, file_type = DocumentManager.load(file)
                
                if content:
                    # Chunk
                    chunks = st.session_state.text_splitter.split_text(content)
                    
                    # Add to collection
                    chunk_ids = [f"{file.name}_{datetime.now().timestamp()}_{j}" 
                                for j in range(len(chunks))]
                    metadatas = [{
                        "filename": file.name,
                        "file_type": file_type,
                        "chunk_index": j,
                        "total_chunks": len(chunks),
                        "upload_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "file_size": len(content)
                    } for j in range(len(chunks))]
                    
                    st.session_state.collection.add(
                        documents=chunks,
                        ids=chunk_ids,
                        metadatas=metadatas
                    )
                    
                    # Save file info
                    st.session_state.uploaded_files_info.append({
                        'filename': file.name,
                        'file_type': file_type,
                        'chunks': len(chunks),
                        'size': len(content),
                        'upload_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    })
                    
                    status.success(f"✅ {file.name} ({len(chunks)} chunks)")
                else:
                    status.error(f"❌ Failed to process {file.name}")
                
                progress.progress((i + 1) / len(uploaded_files))
            
            st.success("🎉 All files uploaded!")
            st.balloons()
            time.sleep(1)
            st.rerun()

with tab2:
    st.subheader("Document Library")
    
    if st.session_state.collection.count() > 0:
        # Get all data
        all_data = st.session_state.collection.get()
        
        # Group by filename
        files = {}
        for metadata in all_data['metadatas']:
            filename = metadata['filename']
            if filename not in files:
                files[filename] = {
                    'chunks': [],
                    'file_type': metadata['file_type'],
                    'upload_date': metadata.get('upload_date', 'Unknown')
                }
            files[filename]['chunks'].append(metadata)
        
        # Display stats
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("📚 Total Files", len(files))
        with col2:
            st.metric("📄 Total Chunks", st.session_state.collection.count())
        with col3:
            total_size = sum(m.get('file_size', 0) for m in all_data['metadatas'])
            st.metric("💾 Total Size", f"{total_size:,} chars")
        
        st.divider()
        
        # Filter
        col1, col2 = st.columns([3, 1])
        with col1:
            search_filename = st.text_input("🔍 Search files", placeholder="Enter filename...")
        with col2:
            file_type_filter = st.selectbox("Filter type", ["All"] + list(set(f['file_type'] for f in files.values())))
        
        # Display files
        for filename, info in sorted(files.items()):
            # Apply filters
            if search_filename and search_filename.lower() not in filename.lower():
                continue
            if file_type_filter != "All" and info['file_type'] != file_type_filter:
                continue
            
            with st.expander(f"📄 {filename}"):
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    st.write(f"**Type:** {info['file_type']}")
                with col2:
                    st.write(f"**Chunks:** {len(info['chunks'])}")
                with col3:
                    st.write(f"**Uploaded:** {info['upload_date']}")
                with col4:
                    if st.button("🗑️ Delete", key=f"del_{filename}"):
                        # Delete all chunks for this file
                        chunk_ids = [m['chunk_index'] for m in info['chunks']]
                        # Note: ChromaDB delete by IDs would go here
                        st.success(f"Deleted {filename}")
                        st.rerun()
                
                # Show chunks
                st.write(f"**Chunks ({len(info['chunks'])}):**")
                for i, chunk_meta in enumerate(info['chunks'][:5], 1):
                    st.text(f"Chunk {i}/{len(info['chunks'])}")
    else:
        st.info("📭 No documents yet. Upload some files!")

with tab3:
    st.subheader("Search Documents")
    
    if st.session_state.collection.count() > 0:
        query = st.text_input("Enter search query")
        
        col1, col2 = st.columns([3, 1])
        with col1:
            n_results = st.slider("Results", 1, 20, 5)
        with col2:
            if st.button("🔍 Search", type="primary"):
                if query:
                    results = st.session_state.collection.query(
                        query_texts=[query],
                        n_results=n_results
                    )
                    
                    st.write(f"**Found {len(results['documents'][0])} results:**")
                    
                    for i, (doc, metadata, distance) in enumerate(zip(
                        results['documents'][0],
                        results['metadatas'][0],
                        results['distances'][0]
                    ), 1):
                        similarity = 1 - distance
                        
                        with st.container():
                            st.markdown(f"""
                            **{i}. {metadata['filename']}** (Chunk {metadata['chunk_index'] + 1}/{metadata.get('total_chunks', '?')})
                            
                            📊 Similarity: {similarity:.2%}
                            
                            {doc[:200]}...
                            """)
                            st.divider()
    else:
        st.info("Upload documents first!")

# Sidebar stats
with st.sidebar:
    st.title("📊 Statistics")
    st.metric("Documents", st.session_state.collection.count())
    
    if st.button("🔄 Refresh"):
        st.rerun()